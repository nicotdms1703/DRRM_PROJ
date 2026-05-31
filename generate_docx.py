import os
from docx import Document
from docx.shared import Pt, Inches

def create_doc():
    doc = Document()
    
    # Title
    doc.add_heading('DBMS PROJECT PROPOSAL', 0)
    
    # 1 - Title Slide
    doc.add_heading('1 — Title Slide', level=1)
    p = doc.add_paragraph()
    p.add_run('Project Title: ').bold = True
    p.add_run('SNSU DRRM (Surigao del Norte State University Disaster Risk Reduction and Management) Resource Management System\n')
    p.add_run('Your Name: ').bold = True
    p.add_run('Nico\n')
    p.add_run('Course / Program: ').bold = True
    p.add_run('BS Information Technology / Computer Science\n')
    p.add_run('Instructor: ').bold = True
    p.add_run('[Instructor\'s Name]\n')
    p.add_run('Date: ').bold = True
    p.add_run('May 28, 2026')
    
    # 2 - Introduction
    doc.add_heading('2 — Introduction', level=1)
    doc.add_paragraph('Background of the system: Disaster Risk Reduction and Management requires prompt, efficient tracking of resources and emergency equipment.', style='List Bullet')
    doc.add_paragraph('Current situation/problem: Current manual processes lead to inaccurate tracking of assets, difficulty monitoring borrowing status, and inadequate predictive capabilities for consumable materials.', style='List Bullet')
    doc.add_paragraph('Importance of database systems: A digitized database system ensures real-time oversight, prevents asset loss, tracks wear and tear, and maintains operational readiness in times of disaster.', style='List Bullet')
    
    # 3 - Problem Statement
    doc.add_heading('3 — Problem Statement', level=1)
    doc.add_paragraph('Lack of real-time visibility into the availability and physical condition of DRRM equipment.', style='List Bullet')
    doc.add_paragraph('Difficulty in tracking borrowed and returned assets, causing delays and possible losses.', style='List Bullet')
    doc.add_paragraph('Lack of predictive alerts for low stock on consumable resources, leading to emergency stockouts.', style='List Bullet')
    doc.add_paragraph('Inefficient logistics tracking for reporting monthly supply deployment and return rates.', style='List Bullet')

    # 4 - Objectives
    doc.add_heading('4 — Objectives', level=1)
    doc.add_paragraph('General Objective: ', style='List Bullet').add_run('To design and develop an AI-powered DRRM Resource Management System that centralizes inventory, tracks borrowing and deployment, and provides predictive analytics.')
    doc.add_paragraph('Specific Objectives:', style='List Bullet')
    doc.add_paragraph('To implement a full CRUD registry for DRRM resources with barcode/QR scanning.', style='List Continue')
    doc.add_paragraph('To design a reliable stock management and borrowing system.', style='List Continue')
    doc.add_paragraph('To integrate AI-driven analytics to forecast consumption and generate restock alerts.', style='List Continue')
    doc.add_paragraph('To improve operational dashboard reporting for real-time decision-making.', style='List Continue')

    # 5 - Scope and Limitations
    doc.add_heading('5 — Scope and Limitations', level=1)
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('Scope: ').bold = True
    p.add_run('The system includes inventory tracking, QR generation and scanning, borrowing/return workflows, stock adjustments, predictive analytics (restock/maintenance alerts), and activity audit trails. Target users are DRRM administrators, officers, and staff.')
    p2 = doc.add_paragraph('', style='List Bullet')
    p2.add_run('Limitations: ').bold = True
    p2.add_run('The system excludes offline synchronization, hardware integration (other than standard barcode scanners/webcams), and external agency data sharing. It does not handle procurement purchasing APIs or financial accounting.')

    # 6 - Proposed System Overview
    doc.add_heading('6 — Proposed System Overview', level=1)
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('Description of your system: ').bold = True
    p.add_run('A web-based application built with Django that manages DRRM assets using a structured relational database to ensure high availability of critical supplies.')
    p2 = doc.add_paragraph('', style='List Bullet')
    p2.add_run('Key features: ').bold = True
    doc.add_paragraph('Real-time Dashboard with KPIs', style='List Continue')
    doc.add_paragraph('QR-based Item Registry', style='List Continue')
    doc.add_paragraph('Stock Transaction Management (In/Out)', style='List Continue')
    doc.add_paragraph('Borrowing tracking system', style='List Continue')
    doc.add_paragraph('AI Alerts and Consumption Analytics', style='List Continue')

    # 7 - Conceptual Framework
    doc.add_heading('7 — Conceptual Framework', level=1)
    doc.add_paragraph('Input -> Process -> Output', style='Intense Quote')
    doc.add_paragraph('Input: User credentials, Asset Details, Borrowing Info, Stock Adjustments, Barcode scans.', style='List Bullet')
    doc.add_paragraph('Process: Authentication, Stock Validation, Predictive Analytics, QR/Barcode Parsing.', style='List Bullet')
    doc.add_paragraph('Output: Dashboard Charts, Restock/Maintenance Alerts, Printable QR Tags, Transaction history.', style='List Bullet')

    # 8 - Database Design
    doc.add_heading('8 — Database Design', level=1)
    doc.add_paragraph('List of tables (Entities):', style='List Bullet')
    doc.add_paragraph('User: System users and authentication.', style='List Continue')
    doc.add_paragraph('Item: The core asset registry.', style='List Continue')
    doc.add_paragraph('StockTransaction: Logs of items coming in and out.', style='List Continue')
    doc.add_paragraph('Borrowing: Records of items lent to personnel and their return status.', style='List Continue')
    doc.add_paragraph('RestockAlert / MaintenanceAlert: AI-generated warnings based on consumption/integrity.', style='List Continue')
    doc.add_paragraph('LogisticsAnalytics: Aggregated monthly performance data.', style='List Continue')
    doc.add_paragraph('ActivityLog: System audit trail.', style='List Continue')
    doc.add_paragraph('Relationships:', style='List Bullet')
    doc.add_paragraph('One-to-Many: Item has many StockTransactions and Borrowings.', style='List Continue')
    doc.add_paragraph('One-to-Many: Item has many RestockAlerts and MaintenanceAlerts.', style='List Continue')
    doc.add_paragraph('One-to-Many: User can have multiple ActivityLogs and acknowledge multiple alerts.', style='List Continue')

    # 9 - ER Diagram
    doc.add_heading('9 — ER Diagram', level=1)
    doc.add_paragraph('The ER diagram features the following core mappings:', style='Normal')
    doc.add_paragraph('USER (1) to (M) ACTIVITY_LOG', style='List Bullet')
    doc.add_paragraph('USER (1) to (M) STOCK_TRANSACTION', style='List Bullet')
    doc.add_paragraph('ITEM (1) to (M) STOCK_TRANSACTION', style='List Bullet')
    doc.add_paragraph('ITEM (1) to (M) BORROWING', style='List Bullet')
    doc.add_paragraph('ITEM (1) to (M) RESTOCK_ALERT', style='List Bullet')
    doc.add_paragraph('Note: For the visual representation, please refer to the Mermaid diagram in the markdown proposal or use standard ER diagram drawing tools mapping these entities.', style='Normal')

    # 10 - Methodology
    doc.add_heading('10 — Methodology', level=1)
    doc.add_paragraph('Development method: Agile Software Development Life Cycle (SDLC)', style='List Bullet')
    doc.add_paragraph('Phases:', style='List Bullet')
    doc.add_paragraph('Planning: Requirement gathering, defining system scope, and analyzing current manual workflows.', style='List Continue')
    doc.add_paragraph('Design: Creating UI/UX mockups, defining the ER diagram, and designing the database schema.', style='List Continue')
    doc.add_paragraph('Development: Building the backend models, implementing the views, integrating AI analytics, and designing the frontend templates.', style='List Continue')
    doc.add_paragraph('Testing: Unit testing for models, testing QR code scanning, usability testing, and fixing edge-case bugs.', style='List Continue')

    # 11 - Tools and Technologies
    doc.add_heading('11 — Tools and Technologies', level=1)
    doc.add_paragraph('Database: SQLite3 (development) / PostgreSQL (production target)', style='List Bullet')
    doc.add_paragraph('Programming Language: Python 3, JavaScript, HTML5, CSS3', style='List Bullet')
    doc.add_paragraph('Tools: Django Web Framework, Bootstrap, Mermaid.js', style='List Bullet')

    # 12 - Expected Results
    doc.add_heading('12 — Expected Results', level=1)
    doc.add_paragraph('Benefits of the system:', style='List Bullet')
    doc.add_paragraph('Faster processing: Equipment can be checked out and returned instantly via QR scanning.', style='List Continue')
    doc.add_paragraph('Better organization: Clear categorizations, structured transaction history, and item integrity tracking.', style='List Continue')
    doc.add_paragraph('Reduced errors: Automatic validation prevents over-borrowing and invalid stock deductions.', style='List Continue')
    doc.add_paragraph('Proactive Management: AI alerts ensure consumables and vital equipment never run out unexpectedly.', style='List Continue')

    # 13 - Timeline
    doc.add_heading('13 — Timeline', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Phase'
    hdr_cells[1].text = 'Activity'
    hdr_cells[2].text = 'Duration'
    
    phases = [
        ('Week 1-2', 'Planning and Requirement Analysis', '2 Weeks'),
        ('Week 3-4', 'Database Design and Wireframing', '2 Weeks'),
        ('Week 5-8', 'Core Development (Backend & Frontend)', '4 Weeks'),
        ('Week 9-10', 'Integration of AI Analytics and QR system', '2 Weeks'),
        ('Week 11', 'System Testing and Debugging', '1 Week'),
        ('Week 12', 'Deployment and Final Presentation', '1 Week')
    ]
    for phase, activity, duration in phases:
        row_cells = table.add_row().cells
        row_cells[0].text = phase
        row_cells[1].text = activity
        row_cells[2].text = duration

    # 14 - Conclusion
    doc.add_heading('14 — Conclusion', level=1)
    doc.add_paragraph('Summary of proposal: This project proposes the creation of the SNSU DRRM Resource Management System, an intelligent web application designed to track emergency response assets, streamline equipment lending, and provide predictive data.', style='List Bullet')
    doc.add_paragraph('Expected impact: By migrating from manual logbooks to a digitized, AI-enhanced database, the DRRM department will ensure full operational readiness, faster response times, and complete accountability for all resources.', style='List Bullet')

    doc.save('PROJECT_PROPOSAL.docx')
    print("Document successfully created: PROJECT_PROPOSAL.docx")

if __name__ == '__main__':
    create_doc()
